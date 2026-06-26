"""
Word document (.docx / .doc) text extraction.
Accepts raw file bytes, returns cleaned text string.

Strategy:
  1. python-docx  — fast, preserves paragraph structure
  2. mammoth      — fallback, good for complex formatting
  3. extract-text — final fallback via CLI
"""

import io
import re
import tempfile
import os


def extract_text_from_docx(file_bytes: bytes, filename: str = "") -> str:
    """Extract plain text from a .docx or .doc file given its raw bytes."""
    ext = os.path.splitext(filename.lower())[1] if filename else ".docx"

    # .doc (legacy binary) → convert to .docx first, then extract
    if ext == ".doc":
        converted = _convert_doc_to_docx(file_bytes)
        if converted:
            file_bytes = converted
        else:
            return ""  # conversion failed

    text = _extract_python_docx(file_bytes)
    if not text.strip():
        text = _extract_mammoth(file_bytes)
    if not text.strip():
        text = _extract_cli(file_bytes)

    return _clean(text)


# ── Extractors ────────────────────────────────────────────────────────────────

def _extract_python_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            line = para.text.strip()
            if line:
                parts.append(line)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_mammoth(file_bytes: bytes) -> str:
    try:
        import mammoth
        result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
        return result.value or ""
    except Exception:
        return ""


def _extract_cli(file_bytes: bytes) -> str:
    """Use extract-text CLI as last resort."""
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        result = os.popen(f"extract-text {tmp_path} 2>/dev/null").read()
        os.unlink(tmp_path)
        return result
    except Exception:
        return ""


def _convert_doc_to_docx(file_bytes: bytes) -> bytes | None:
    """Convert legacy .doc bytes to .docx bytes via LibreOffice."""
    try:
        import subprocess, glob
        with tempfile.TemporaryDirectory() as tmpdir:
            src = os.path.join(tmpdir, "input.doc")
            with open(src, "wb") as f:
                f.write(file_bytes)
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx", "--outdir", tmpdir, src],
                capture_output=True, timeout=30
            )
            matches = glob.glob(os.path.join(tmpdir, "*.docx"))
            if matches:
                with open(matches[0], "rb") as f:
                    return f.read()
    except Exception:
        pass
    return None


# ── Text cleaner ──────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
